import io
import zipfile
from datetime import date

from docx import Document

from scripts import twinning_monitor as tm


LISTING_HTML = """
<table><tbody>
<tr><td>
  <span class="label label-default">2026-07-09</span>
  <span class="label label-info">Gambia</span>
  <span class="label label-primary">Rolnictwo</span><br>
  <a href="/fiszki-twinning/example-one/">Example One</a>
</td></tr>
<tr><td>
  <span class="label label-default">2026-06-22</span>
  <span class="label label-info">Macedonia Północna</span>
  <span class="label label-primary">Sprawiedliwość</span><br>
  <a href="/fiszki-twinning/example-two/">Example Two</a>
</td></tr>
</tbody></table>
"""


DETAIL_HTML = """
<div class="container-fluid">
  <div class="row"><div>Beneficjent:</div><div>Gambia</div></div>
  <div class="row"><div>Numer:</div><div>GM 24 NDICI AG 01 26</div></div>
  <div class="row"><div>Tytuł:</div><div>Food safety support</div></div>
  <div class="row"><div>Obszar:</div><div>Rolnictwo</div></div>
</div>
Załącznik: <a href=/media/example.zip>Pobierz</a>
<p>Termin przesłania oferty bliźniaczej do MSZ: <strong>29 września 2026 r.</strong></p>
<p>Termin przesłania oferty przez MSZ do beneficjenta: 1 października 2026 r.</p>
"""


def test_parse_listing_extracts_labels_and_absolute_links():
    offers = tm.parse_listing(LISTING_HTML)
    assert len(offers) == 2
    assert offers[0].country == "Gambia"
    assert offers[0].url == "https://twinning.msz.gov.pl/fiszki-twinning/example-one/"
    assert offers[1].area == "Sprawiedliwość"


def test_ascii_url_encodes_unicode_slug_without_double_encoding():
    value = tm._ascii_url("https://example.test/fiszki/à-côte/?q=już%20ok")
    assert value == "https://example.test/fiszki/%C3%A0-c%C3%B4te/?q=ju%C5%BC%20ok"


def test_parse_detail_extracts_core_fields_and_deadlines():
    detail = tm.parse_detail(DETAIL_HTML, "https://twinning.msz.gov.pl/fiszki-twinning/example-one/")
    assert detail["beneficiary"] == "Gambia"
    assert detail["reference"] == "GM 24 NDICI AG 01 26"
    assert detail["attachmentUrl"] == "https://twinning.msz.gov.pl/media/example.zip"
    assert tm.parse_polish_date(detail["mszDeadline"]) == date(2026, 9, 29)


def test_list_active_offers_skips_expired_and_is_deterministic(monkeypatch):
    expired = DETAIL_HTML.replace("29 września 2026", "29 lipca 2026")

    def fake_fetch(url):
        if url.endswith("fiszki-twinning"):
            return LISTING_HTML
        return DETAIL_HTML if url.endswith("example-one/") else expired

    monkeypatch.setattr(tm, "fetch_text", fake_fetch)
    monkeypatch.setattr(tm, "_attachment_metadata", lambda url: {"url": url, "etag": "abc"})
    result = tm.list_active_offers(today=date(2026, 8, 10), lookback_days=180)
    assert [offer["offerId"] for offer in result["activeOffers"]] == ["GM 24 NDICI AG 01 26"]
    assert len(result["activeOffers"][0]["contentHash"]) == 64


def test_extract_offer_prefers_annex_c1_docx(monkeypatch):
    primary = Document()
    primary.add_paragraph("ANNEX C1: Twinning Fiche")
    primary.add_paragraph("3.6 Means/input from the EU Member State Partner Administration")
    primary.add_paragraph("Profile and tasks of the RTA: ten years of relevant experience")
    primary_buffer = io.BytesIO()
    primary.save(primary_buffer)

    other = Document()
    other.add_paragraph("Supporting strategy")
    other_buffer = io.BytesIO()
    other.save(other_buffer)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as archive:
        archive.writestr("Annex C1.docx", primary_buffer.getvalue())
        archive.writestr("Strategy.docx", other_buffer.getvalue())

    monkeypatch.setattr(tm, "fetch_text", lambda url: DETAIL_HTML)
    monkeypatch.setattr(
        tm,
        "_request",
        lambda url, timeout=30, method="GET": (zip_buffer.getvalue(), {}, "https://example.test/example.zip"),
    )
    result = tm.extract_offer("https://twinning.msz.gov.pl/fiszki-twinning/example-one/")
    assert result["primaryDocument"] == "Annex C1.docx"
    assert "Profile and tasks of the RTA" in result["analysisText"]


def test_analysis_excerpt_keeps_requirements_from_long_document():
    text = "Introduction " + ("background " * 14_000) + " Profile and tasks of the RTA: minimum 10 years."
    excerpt = tm._analysis_excerpt(text)
    assert len(excerpt) <= tm.MAX_ANALYSIS_CHARS
    assert "minimum 10 years" in excerpt


def test_durable_notification_state_marks_new_sent_and_updated(tmp_path):
    state_path = tmp_path / "twinning-state.json"
    base = {"activeOffers": [{"offerId": "TEST-1", "contentHash": "hash-one"}]}
    assert tm.apply_notification_state(base, state_path)["activeOffers"][0]["notificationType"] == "new"

    saved = tm.acknowledge_offer(
        "TEST-1",
        "hash-one",
        ["one@example.test", "two@example.test"],
        "resend-123",
        path=state_path,
    )
    assert saved["resendId"] == "resend-123"
    assert tm.apply_notification_state(base, state_path)["activeOffers"][0]["notificationType"] is None

    changed = {"activeOffers": [{"offerId": "TEST-1", "contentHash": "hash-two"}]}
    assert tm.apply_notification_state(changed, state_path)["activeOffers"][0]["notificationType"] == "updated"
    assert tm.load_notification_state(state_path)["offers"]["TEST-1"]["recipients"] == [
        "one@example.test",
        "two@example.test",
    ]


def test_nonmatching_offer_moves_through_digest_queue(tmp_path):
    state_path = tmp_path / "twinning-state.json"
    item = {
        "offerId": "TEST-2",
        "contentHash": "hash-two",
        "title": "Agriculture controls",
        "country": "Exampleland",
        "area": "Rolnictwo",
        "mszDeadline": "20 sierpnia 2026 r.",
        "url": "https://example.test/offer",
        "fitBand": "no_fit",
        "fitScore": 5,
        "fitReason": "Dziedzina poza profilem.",
        "bestEntryRole": "STE rolnictwa",
    }
    queued = tm.queue_digest_offer(item, path=state_path)
    assert queued["status"] == "digest_pending"
    assert tm.pending_digest_offers(state_path) == [item]
    assert tm.apply_notification_state(
        {"activeOffers": [{"offerId": "TEST-2", "contentHash": "hash-two"}]},
        state_path,
    )["activeOffers"][0]["notificationType"] is None

    changed = tm.acknowledge_digest(
        ["TEST-2"],
        ["one@example.test", "two@example.test"],
        "digest-resend-1",
        path=state_path,
    )
    assert changed == 1
    assert tm.pending_digest_offers(state_path) == []
    entry = tm.load_notification_state(state_path)["offers"]["TEST-2"]
    assert entry["status"] == "digest_sent"
    assert entry["resendId"] == "digest-resend-1"
