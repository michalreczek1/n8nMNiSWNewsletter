import io
import http.client
import json
import re
import socket
import ssl
import sys
from html import unescape
from urllib.parse import urljoin
from urllib.request import HTTPHandler, HTTPSHandler, Request, build_opener

from docx import Document


HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "pl-PL,pl;q=0.9,en-US;q=0.8,en;q=0.7",
    "Cache-Control": "no-cache",
}


class IPv4HTTPConnection(http.client.HTTPConnection):
    def connect(self):
        addrs = socket.getaddrinfo(self.host, self.port, socket.AF_INET, socket.SOCK_STREAM)
        last_error = None
        for family, socktype, proto, _, sockaddr in addrs:
            sock = socket.socket(family, socktype, proto)
            sock.settimeout(self.timeout)
            try:
                sock.connect(sockaddr)
                self.sock = sock
                return
            except OSError as exc:
                last_error = exc
                sock.close()
        if last_error is not None:
            raise last_error
        raise OSError(f"Could not resolve IPv4 address for {self.host}")


class IPv4HTTPSConnection(http.client.HTTPSConnection):
    def connect(self):
        addrs = socket.getaddrinfo(self.host, self.port, socket.AF_INET, socket.SOCK_STREAM)
        last_error = None
        for family, socktype, proto, _, sockaddr in addrs:
            sock = socket.socket(family, socktype, proto)
            sock.settimeout(self.timeout)
            try:
                sock.connect(sockaddr)
                if self._tunnel_host:
                    self.sock = sock
                    self._tunnel()
                    sock = self.sock
                self.sock = self._context.wrap_socket(sock, server_hostname=self.host)
                return
            except OSError as exc:
                last_error = exc
                sock.close()
        if last_error is not None:
            raise last_error
        raise OSError(f"Could not resolve IPv4 address for {self.host}")


class IPv4HTTPHandler(HTTPHandler):
    def http_open(self, req):
        return self.do_open(IPv4HTTPConnection, req)


class IPv4HTTPSHandler(HTTPSHandler):
    def https_open(self, req):
        return self.do_open(IPv4HTTPSConnection, req)


URL_OPENER = build_opener(IPv4HTTPHandler, IPv4HTTPSHandler(context=ssl.create_default_context()))


def fetch_bytes(url: str, timeout: int = 60):
    req = Request(url, headers=HEADERS)
    with URL_OPENER.open(req, timeout=timeout) as resp:
        return resp.read(), resp.headers.get_content_type() or "", resp.geturl()


def fetch_html(url: str) -> str:
    data, _, _ = fetch_bytes(url)
    for encoding in ("utf-8", "cp1250", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def clean_html(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", " ")
    text = text.replace("\xa0", " ")
    text = unescape(text)
    text = re.sub(r"[\u0000-\u0008\u000b-\u001f]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s*\r?\n\s*", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def dedupe_text_blocks(blocks):
    seen = set()
    out = []
    for block in blocks:
        text = clean_text(block)
        if not text:
            continue
        key = re.sub(r"\s+", " ", text).strip().lower()
        if len(key) > 500:
            key = key[:500]
        if key in seen:
            continue
        seen.add(key)
        out.append(text)
    return out


def sanitize_fallback_summary_for_ai(text: str) -> str:
    summary = clean_text(text)
    if not summary:
        return ""
    lowered = summary.lower()
    bad_starts = (
        "projekt dotyczy:",
        "projektu nie dotyczy",
        "nie udało się pobrać",
    )
    if lowered.startswith(bad_starts):
        return ""
    if len(summary) < 40:
        return ""
    return summary


def extract_label(html: str, label: str) -> str:
    patterns = [
        rf"{re.escape(label)}:</div>\s*<div[^>]*>([\s\S]{{0,1200}}?)</div>",
        rf"{re.escape(label)}[\s\S]{{0,120}}?<td[^>]*>([\s\S]{{0,1200}}?)</td>",
        rf"{re.escape(label)}[\s\S]{{0,120}}?<dd[^>]*>([\s\S]{{0,1200}}?)</dd>",
    ]
    for pattern in patterns:
        match = re.search(pattern, html, flags=re.I)
        if match:
            return clean_html(match.group(1))
    return ""


def parse_stage_links(html: str):
    out = []
    pattern = re.compile(
        r'<a href="(/projekt/\d+/katalog/(\d+)#\2)">\s*([^<]+?)\s*</a>\s*.*?Data ostatniej modyfikacji:\s*([0-9]{2}-[0-9]{2}-[0-9]{4})',
        re.I | re.S,
    )
    for match in pattern.finditer(html):
        out.append(
            {
                "url": urljoin("https://legislacja.gov.pl", match.group(1)),
                "catalogId": match.group(2),
                "label": clean_text(match.group(3)),
                "modified": match.group(4),
                "pos": match.start(),
            }
        )
    return out


def parse_doc_links(html: str):
    out = []
    pattern = re.compile(
        r'<a href="(/docs/[^"]+/dokument[^"]+)">([\s\S]*?)</a>',
        re.I | re.S,
    )
    for match in pattern.finditer(html):
        label = clean_html(match.group(2))
        if not label:
            continue
        out.append(
            {
                "url": urljoin("https://legislacja.gov.pl", match.group(1)),
                "label": label,
                "pos": match.start(),
            }
        )
    return out


def choose_doc(doc_links):
    priorities = ("Uzasadnienie", "Ocena skutków regulacji", "OSR", "Projekt")
    for label in priorities:
        preferred = [d for d in doc_links if d["label"].lower() == label.lower()]
        if preferred:
            return preferred[-1]
    excluded = (
        "rozdzielnik",
        "pismo",
        "tabela uwag",
        "uwagi",
        "załączniki",
    )
    filtered = [
        d
        for d in doc_links
        if not any(marker in d["label"].lower() for marker in excluded)
    ]
    return filtered[-1] if filtered else None


def find_stage_for_doc(stage_links, doc_pos: int) -> str:
    prior = [s for s in stage_links if s["pos"] < doc_pos]
    if prior:
        return prior[-1]["label"]
    return stage_links[-1]["label"] if stage_links else ""


def extract_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    blocks = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                if cell_text:
                    blocks.extend(part.strip() for part in cell_text.splitlines() if part.strip())
    deduped = dedupe_text_blocks(blocks)
    return clean_text("\n".join(deduped))


def extract_doc_text(data: bytes) -> str:
    decoded = data.decode("utf-16le", errors="ignore")
    decoded = clean_text(decoded)
    markers = ["UZASADNIENIE", "Uzasadnienie", "Projekt ", "Ocena Skutków Regulacji", "OSR"]
    starts = [decoded.find(marker) for marker in markers if marker in decoded]
    if starts:
        decoded = decoded[min(starts):]
    decoded = re.split(r"(?:[\u0800-\uffff]\s*){4,}", decoded, maxsplit=1)[0]
    lines = []
    for line in decoded.splitlines():
        line = line.strip()
        if len(line) < 3:
            continue
        line = re.sub(r"[\u0800-\uffff]+", " ", line)
        line = clean_text(line)
        if sum(ch.isalpha() for ch in line) < 3:
            continue
        lines.append(line)
    return clean_text("\n".join(lines))


def extract_document_text(url: str):
    data, content_type, final_url = fetch_bytes(url)
    lowered_url = final_url.lower()
    if lowered_url.endswith(".docx") or "wordprocessingml.document" in content_type:
        text = extract_docx_text(data)
    elif lowered_url.endswith(".doc") or "application/msword" in content_type:
        text = extract_doc_text(data)
    else:
        text = clean_text(data.decode("utf-8", errors="ignore"))
    return text, content_type, final_url


def summarize(text: str, title: str) -> str:
    source = clean_text(text).replace("\n", " ")
    if not source:
        return f"Projekt dotyczy: {title}."
    for marker in (
        "Rekomendowane rozwiązanie, w tym planowane narzędzia interwencji i oczekiwany efekt",
        "Rekomendowane rozwiązanie",
        "Jaki problem jest rozwiązywany?",
        "Projektowane regulacje",
        "Istota rozwiązań",
    ):
        if marker in source:
            source = source[source.find(marker):]
            break
    source = re.sub(r"^UZASADNIENIE\s*", "", source, flags=re.I)
    source = re.sub(r"^(Rekomendowane rozwiązanie[^A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż0-9]*)", "", source, flags=re.I)
    source = re.sub(r"^(Jaki problem jest rozwiązywany\?[^A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż0-9]*)", "", source, flags=re.I)
    source = re.sub(r"\bart\.\s*,", "art.", source, flags=re.I)
    source = re.sub(r"\s{2,}", " ", source).strip()
    normalized_title = clean_text(title)
    if normalized_title and source.lower().startswith(normalized_title.lower()):
        source = source[len(normalized_title):].lstrip(" :-–—\n")
    direct_patterns = [
        r"(Celem projektu[^.!?]{0,900}[.!?])",
        r"(Projekt[^.!?]{0,900}(?:wprowadza|określa|zmienia|reguluje)[^.!?]{0,900}[.!?])",
        r"(Przedmiotem projektu[^.!?]{0,900}[.!?])",
        r"(Projekt zakłada:[^0-9]{0,900})",
        r"(Projekt zakłada[^.!?]{0,900}[.!?])",
        r"(Projekt przewiduje[^.!?]{0,900}[.!?])",
        r"(Zmiana warunku[^.!?]{0,900}[.!?])",
        r"(Projektowane regulacje[^.!?]{0,900}[.!?])",
        r"(Regulacja[^.!?]{0,900}(?:polega|zakłada|umożliwia|rozszerza)[^.!?]{0,900}[.!?])",
    ]
    for pattern in direct_patterns:
        match = re.search(pattern, source, flags=re.I)
        if match:
            candidate = clean_text(match.group(1))
            candidate = re.split(r"\s+2\)\s+", candidate, maxsplit=1)[0]
            return candidate[:600]
    sentences = re.findall(r"[^.!?]{40,350}[.!?]", source)
    if sentences:
        return clean_text(" ".join(sentences[:2]))[:600]
    return source[:600]


def build_ai_context(title: str, applicant: str, stage_label: str, project_type: str, number: str, date_created: str, date_updated: str, document_label: str, chosen_text: str, fallback_summary: str) -> str:
    source = clean_text(chosen_text).replace("\n", " ")
    source = re.sub(r"\s{2,}", " ", source).strip()
    source = source[:7000]
    blocks = [
        f"Tytul referencyjny projektu (nie powtarzaj go doslownie w summary): {clean_text(title)}",
        f"Autor / wnioskodawca: {clean_text(applicant)}",
        f"Typ projektu: {clean_text(project_type)}",
        f"Numer projektu: {clean_text(number)}",
        f"Etap: {clean_text(stage_label)}",
        f"Data utworzenia: {clean_text(date_created)}",
        f"Data aktualizacji: {clean_text(date_updated)}",
        f"Zrodlo tresci: {clean_text(document_label)}",
        f"Heurystyczny skrot zmian: {sanitize_fallback_summary_for_ai(fallback_summary)}",
        "Zadanie: opisz autora projektu i 2-4 najwazniejsze zmiany lub rozwiazania. Nie parafrazuj samego tytulu aktu.",
        f"Tresc dokumentu lub projektu: {source}",
    ]
    return "\n".join(block for block in blocks if block and not block.endswith(": "))


def infer_project_type(title: str) -> str:
    lower = title.lower()
    if lower.startswith("projekt ustawy"):
        return "Projekt ustawy"
    if lower.startswith("projekt rozporządzenia"):
        return "Projekt rozporządzenia"
    if lower.startswith("projekt uchwały"):
        return "Projekt uchwały"
    return ""


def extract_project(project_url: str, from_email: str, to_email: str, days_lookback: int = 14, score_threshold: float = 0):
    try:
        html = fetch_html(project_url)
        if "Request Rejected" in html or "Żądanie nie może zostać zrealizowane" in html:
            raise RuntimeError("RCL odrzucił żądanie do projektu.")

        title_match = re.search(r'<div class="rcl-title">\s*([\s\S]{1,1200}?)</div>', html, flags=re.I)
        title = clean_html(title_match.group(1)) if title_match else ""
        applicant = extract_label(html, "Wnioskodawca") or extract_label(html, "Organ wnioskujący") or extract_label(html, "Podmiot odpowiedzialny")
        date_created = extract_label(html, "Data utworzenia")
        status_project = extract_label(html, "Status projektu")
        number = extract_label(html, "Numer z wykazu") or extract_label(html, "Numer projektu")
        date_updated = extract_label(html, "Data aktualizacji")
        stage_links = parse_stage_links(html)
        doc_links = parse_doc_links(html)

        if not doc_links:
            for stage in reversed(stage_links):
                stage_html = fetch_html(stage["url"])
                stage_docs = parse_doc_links(stage_html)
                if stage_docs:
                    doc_links = stage_docs
                    html = stage_html
                    break

        chosen_doc = choose_doc(doc_links)
        chosen_text = ""
        chosen_doc_url = ""
        chosen_doc_label = ""
        stage_label = ""
        content_type = ""

        if chosen_doc:
            chosen_doc_url = chosen_doc["url"]
            chosen_doc_label = chosen_doc["label"]
            stage_label = find_stage_for_doc(stage_links, chosen_doc["pos"])
            chosen_text, content_type, chosen_doc_url = extract_document_text(chosen_doc_url)

        if not stage_label and stage_links:
            stage_label = stage_links[-1]["label"]

        summary = summarize(chosen_text, title or project_url)
        ai_context = build_ai_context(
            title=title or project_url,
            applicant=applicant,
            stage_label=stage_label or status_project,
            project_type=infer_project_type(title),
            number=number,
            date_created=date_created,
            date_updated=date_updated or (stage_links[-1]["modified"] if stage_links else ""),
            document_label=chosen_doc_label,
            chosen_text=chosen_text,
            fallback_summary=summary,
        )
        analysis_text = clean_text(" ".join(filter(None, [title, applicant, stage_label, status_project, chosen_text])))[:14000]

        payload = {
            "projectId": re.search(r"/projekt/(\d+)", project_url).group(1) if re.search(r"/projekt/(\d+)", project_url) else project_url,
            "title": title or project_url,
            "applicant": applicant,
            "stage": stage_label or status_project,
            "number": number,
            "projectType": infer_project_type(title),
            "dateCreated": date_created,
            "dateUpdated": date_updated or (stage_links[-1]["modified"] if stage_links else ""),
            "summary": summary,
            "fallbackSummary": summary,
            "analysisText": analysis_text,
            "aiContext": ai_context,
            "documentUrl": chosen_doc_url,
            "documentLabel": chosen_doc_label,
            "documentContentType": content_type,
            "url": project_url,
            "fromEmail": from_email,
            "toEmail": to_email,
            "daysLookback": days_lookback,
            "scoreThreshold": score_threshold,
            "summarySource": "justification-doc" if chosen_text else "project-page",
        }
        return payload
    except Exception as exc:
        payload = {
            "error": True,
            "message": str(exc),
            "projectId": re.search(r"/projekt/(\d+)", project_url).group(1) if re.search(r"/projekt/(\d+)", project_url) else project_url,
            "title": project_url,
            "applicant": "",
            "stage": "",
            "number": "",
            "projectType": "",
            "dateCreated": "",
            "dateUpdated": "",
            "summary": f"Nie udało się pobrać uzasadnienia projektu: {project_url}",
            "fallbackSummary": f"Nie udało się pobrać uzasadnienia projektu: {project_url}",
            "analysisText": "",
            "aiContext": "",
            "documentUrl": "",
            "documentLabel": "",
            "documentContentType": "",
            "url": project_url,
            "fromEmail": from_email,
            "toEmail": to_email,
            "daysLookback": days_lookback,
            "scoreThreshold": score_threshold,
            "summarySource": "error",
        }
        return payload


def main():
    project_url = sys.argv[1]
    from_email = sys.argv[2]
    to_email = sys.argv[3]
    days_lookback = int(sys.argv[4]) if len(sys.argv) > 4 and str(sys.argv[4]).strip() else 14
    score_threshold = float(sys.argv[5]) if len(sys.argv) > 5 and str(sys.argv[5]).strip() else 0
    payload = extract_project(project_url, from_email, to_email, days_lookback, score_threshold)
    print(json.dumps(payload, ensure_ascii=True))


if __name__ == "__main__":
    main()
