"""Fetch, extract and track active Twinning fiches published by the Polish MFA."""

from __future__ import annotations

import hashlib
import io
import json
import os
import re
import tempfile
import threading
import time
import zipfile
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable
from urllib.error import HTTPError, URLError
from urllib.parse import quote, urljoin, urlsplit, urlunsplit
from urllib.request import Request, urlopen

from docx import Document
from pypdf import PdfReader


BASE_URL = "https://twinning.msz.gov.pl"
LIST_URL = f"{BASE_URL}/fiszki-twinning"
USER_AGENT = "FamilyOS-Twinning-Monitor/1.0 (+https://familyos.pl)"
MAX_DOWNLOAD_BYTES = 40 * 1024 * 1024
MAX_DOCUMENT_BYTES = 25 * 1024 * 1024
MAX_ANALYSIS_CHARS = 40_000
DEFAULT_STATE_PATH = Path(os.getenv("TWINNING_STATE_PATH", "/opt/n8n-app/data/twinning-state.json"))
_STATE_LOCK = threading.RLock()


POLISH_MONTHS = {
    "stycznia": 1,
    "lutego": 2,
    "marca": 3,
    "kwietnia": 4,
    "maja": 5,
    "czerwca": 6,
    "lipca": 7,
    "sierpnia": 8,
    "września": 9,
    "wrzesnia": 9,
    "października": 10,
    "pazdziernika": 10,
    "listopada": 11,
    "grudnia": 12,
}


class TwinningError(RuntimeError):
    """A recoverable source or document-processing error."""


def load_notification_state(path: str | Path = DEFAULT_STATE_PATH) -> dict:
    state_path = Path(path)
    if not state_path.exists():
        return {"offers": {}}
    try:
        payload = json.loads(state_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise TwinningError(f"Nie można odczytać rejestru wysyłek {state_path}: {exc}") from exc
    if not isinstance(payload, dict) or not isinstance(payload.get("offers"), dict):
        raise TwinningError(f"Nieprawidłowy format rejestru wysyłek: {state_path}")
    return payload


def apply_notification_state(payload: dict, path: str | Path = DEFAULT_STATE_PATH) -> dict:
    state = load_notification_state(path)
    sent_offers = state["offers"]
    for offer in payload.get("activeOffers", []):
        previous = sent_offers.get(offer.get("offerId"))
        if previous and previous.get("contentHash") == offer.get("contentHash"):
            offer["notificationType"] = None
        else:
            offer["notificationType"] = "updated" if previous else "new"
    return payload


def _write_notification_state(state: dict, state_path: Path) -> None:
    state_path.parent.mkdir(parents=True, exist_ok=True)
    temp_name = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            dir=state_path.parent,
            prefix=f".{state_path.name}.",
            suffix=".tmp",
            delete=False,
        ) as handle:
            temp_name = handle.name
            json.dump(state, handle, ensure_ascii=False, indent=2, sort_keys=True)
            handle.write("\n")
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temp_name, state_path)
    finally:
        if temp_name and os.path.exists(temp_name):
            os.unlink(temp_name)


def acknowledge_offer(
    offer_id: str,
    content_hash: str,
    recipients: list[str] | None = None,
    resend_id: str | None = None,
    *,
    path: str | Path = DEFAULT_STATE_PATH,
) -> dict:
    if not offer_id or not content_hash:
        raise ValueError("offerId and contentHash are required")
    state_path = Path(path)
    with _STATE_LOCK:
        state = load_notification_state(state_path)
        entry = {
            "contentHash": content_hash,
            "status": "immediate_sent",
            "sentAt": datetime.now().astimezone().isoformat(timespec="seconds"),
            "recipients": recipients or [],
            "resendId": resend_id or None,
        }
        state["offers"][offer_id] = entry
        _write_notification_state(state, state_path)
    return entry


def queue_digest_offer(item: dict, *, path: str | Path = DEFAULT_STATE_PATH) -> dict:
    required = ("offerId", "contentHash", "title", "url")
    if any(not item.get(key) for key in required):
        raise ValueError("offerId, contentHash, title and url are required")
    state_path = Path(path)
    with _STATE_LOCK:
        state = load_notification_state(state_path)
        entry = {
            "contentHash": str(item["contentHash"]),
            "status": "digest_pending",
            "queuedAt": datetime.now().astimezone().isoformat(timespec="seconds"),
            "digestItem": {
                key: item.get(key)
                for key in (
                    "offerId",
                    "contentHash",
                    "title",
                    "country",
                    "area",
                    "mszDeadline",
                    "url",
                    "fitBand",
                    "fitScore",
                    "fitReason",
                    "bestEntryRole",
                )
            },
        }
        state["offers"][str(item["offerId"])] = entry
        _write_notification_state(state, state_path)
    return entry


def pending_digest_offers(path: str | Path = DEFAULT_STATE_PATH) -> list[dict]:
    state = load_notification_state(path)
    return [
        entry["digestItem"]
        for entry in state["offers"].values()
        if entry.get("status") == "digest_pending" and isinstance(entry.get("digestItem"), dict)
    ]


def acknowledge_digest(
    offer_ids: list[str],
    recipients: list[str] | None = None,
    resend_id: str | None = None,
    *,
    path: str | Path = DEFAULT_STATE_PATH,
) -> int:
    state_path = Path(path)
    with _STATE_LOCK:
        state = load_notification_state(state_path)
        sent_at = datetime.now().astimezone().isoformat(timespec="seconds")
        changed = 0
        for offer_id in offer_ids:
            entry = state["offers"].get(offer_id)
            if not entry or entry.get("status") != "digest_pending":
                continue
            entry.update(
                {
                    "status": "digest_sent",
                    "digestSentAt": sent_at,
                    "recipients": recipients or [],
                    "resendId": resend_id or None,
                }
            )
            changed += 1
        if changed:
            _write_notification_state(state, state_path)
    return changed


def _clean(value: str | None) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _ascii_url(url: str) -> str:
    parts = urlsplit(url)
    return urlunsplit(
        (
            parts.scheme,
            parts.netloc.encode("idna").decode("ascii"),
            quote(parts.path, safe="/%:@!$&'()*+,;=-._~"),
            quote(parts.query, safe="=&%:@!$'()*+,;/?-._~"),
            quote(parts.fragment, safe="-._~"),
        )
    )


def _request(url: str, *, timeout: int = 30, method: str = "GET") -> tuple[bytes, dict[str, str], str]:
    request = Request(_ascii_url(url), headers={"User-Agent": USER_AGENT}, method=method)
    last_error: Exception | None = None
    for attempt in range(3):
        try:
            with urlopen(request, timeout=timeout) as response:
                length = int(response.headers.get("Content-Length") or 0)
                if length > MAX_DOWNLOAD_BYTES:
                    raise TwinningError(f"Plik przekracza limit {MAX_DOWNLOAD_BYTES} B: {url}")
                data = b"" if method == "HEAD" else response.read(MAX_DOWNLOAD_BYTES + 1)
                if len(data) > MAX_DOWNLOAD_BYTES:
                    raise TwinningError(f"Pobrano plik przekraczający limit: {url}")
                headers = {key.lower(): value for key, value in response.headers.items()}
                return data, headers, response.geturl()
        except (HTTPError, URLError, TimeoutError, OSError) as exc:
            last_error = exc
            if attempt < 2:
                time.sleep(0.5 * (attempt + 1))
    raise TwinningError(f"Nie udało się pobrać {url}: {last_error}")


def fetch_text(url: str) -> str:
    data, headers, _ = _request(url)
    charset_match = re.search(r"charset=([^;\s]+)", headers.get("content-type", ""), re.I)
    encodings = [charset_match.group(1).strip('"\'')] if charset_match else []
    encodings.extend(["utf-8", "cp1250", "latin-1"])
    for encoding in encodings:
        try:
            return data.decode(encoding)
        except (LookupError, UnicodeDecodeError):
            continue
    return data.decode("utf-8", errors="replace")


@dataclass
class ListingOffer:
    published_date: str
    country: str
    area: str
    title: str
    url: str


class _ListingParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.in_row = False
        self.row_depth = 0
        self.current_label: str | None = None
        self.current_link: str | None = None
        self.labels: list[str] = []
        self.label_parts: list[str] = []
        self.link_parts: list[str] = []
        self.offers: list[ListingOffer] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attr = dict(attrs)
        if tag == "tr":
            if self.in_row:
                self.row_depth += 1
            else:
                self.in_row = True
                self.row_depth = 1
                self.labels = []
                self.current_link = None
                self.link_parts = []
        if not self.in_row:
            return
        if tag == "span" and "label" in (attr.get("class") or "").split():
            self.current_label = attr.get("class") or ""
            self.label_parts = []
        if tag == "a":
            href = attr.get("href") or ""
            if "/fiszki-twinning/" in href and href.rstrip("/") != "/fiszki-twinning":
                self.current_link = href
                self.link_parts = []

    def handle_data(self, data: str) -> None:
        if self.current_label is not None:
            self.label_parts.append(data)
        if self.current_link is not None:
            self.link_parts.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag == "span" and self.current_label is not None:
            value = _clean("".join(self.label_parts))
            if value:
                self.labels.append(value)
            self.current_label = None
            self.label_parts = []
        if tag == "a" and self.current_link is not None:
            # Link data is retained until the row closes.
            pass
        if tag == "tr" and self.in_row:
            self.row_depth -= 1
            if self.row_depth == 0:
                title = _clean("".join(self.link_parts))
                if self.current_link and title and len(self.labels) >= 3:
                    self.offers.append(
                        ListingOffer(
                            published_date=self.labels[0],
                            country=self.labels[1],
                            area=self.labels[2],
                            title=title,
                            url=urljoin(BASE_URL, self.current_link),
                        )
                    )
                self.in_row = False
                self.current_link = None


def parse_listing(html: str) -> list[ListingOffer]:
    parser = _ListingParser()
    parser.feed(html)
    seen: set[str] = set()
    result: list[ListingOffer] = []
    for offer in parser.offers:
        if offer.url in seen:
            continue
        seen.add(offer.url)
        result.append(offer)
    return result


class _DetailParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.depth = 0
        self.in_row = False
        self.row_depth = 0
        self.row_parts: list[str] = []
        self.rows: list[str] = []
        self.text_parts: list[str] = []
        self.attachment_url = ""

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attr = dict(attrs)
        if tag == "div" and "row" in (attr.get("class") or "").split():
            if not self.in_row:
                self.in_row = True
                self.row_depth = 1
                self.row_parts = []
            else:
                self.row_depth += 1
        elif tag == "div" and self.in_row:
            self.row_depth += 1
        if tag == "a":
            href = attr.get("href") or ""
            if href.lower().split("?")[0].endswith((".zip", ".pdf", ".docx", ".doc")):
                self.attachment_url = urljoin(BASE_URL, href)
        if tag in {"br", "p", "li", "div", "h1", "h2", "h3"}:
            self.text_parts.append("\n")

    def handle_data(self, data: str) -> None:
        self.text_parts.append(data)
        if self.in_row:
            self.row_parts.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag == "div" and self.in_row:
            self.row_depth -= 1
            if self.row_depth == 0:
                value = _clean(" ".join(self.row_parts))
                if value:
                    self.rows.append(value)
                self.in_row = False
                self.row_parts = []
        if tag in {"p", "li", "div", "h1", "h2", "h3"}:
            self.text_parts.append("\n")


def parse_polish_date(value: str | None) -> date | None:
    text = _clean(value).lower().replace("r.", "")
    match = re.search(r"\b(\d{1,2})\s+([a-ząćęłńóśźż]+)\s+(20\d{2})\b", text)
    if match:
        month = POLISH_MONTHS.get(match.group(2))
        if month:
            return date(int(match.group(3)), month, int(match.group(1)))
    for pattern in (r"\b(20\d{2})-(\d{2})-(\d{2})\b", r"\b(\d{2})[./-](\d{2})[./-](20\d{2})\b"):
        match = re.search(pattern, text)
        if not match:
            continue
        try:
            if match.group(1).startswith("20"):
                return date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
            return date(int(match.group(3)), int(match.group(2)), int(match.group(1)))
        except ValueError:
            pass
    return None


def parse_detail(html: str, source_url: str) -> dict[str, str | None]:
    parser = _DetailParser()
    parser.feed(html)
    fields: dict[str, str | None] = {
        "beneficiary": None,
        "reference": None,
        "title": None,
        "area": None,
    }
    labels = {
        "beneficjent": "beneficiary",
        "numer": "reference",
        "tytuł": "title",
        "obszar": "area",
    }
    for row in parser.rows:
        match = re.match(r"^(Beneficjent|Numer|Tytuł|Obszar)\s*:\s*(.+)$", row, re.I)
        if match:
            fields[labels[match.group(1).lower()]] = _clean(match.group(2))

    page_text = "\n".join(_clean(part) for part in "".join(parser.text_parts).splitlines() if _clean(part))
    content_start = page_text.find("Beneficjent:")
    if content_start >= 0:
        page_text = page_text[content_start:]
    fallback_patterns = {
        "beneficiary": r"(?:^|\n)Beneficjent:\s*\n?([^\n]+)",
        "reference": r"(?:^|\n)Numer:\s*\n?([^\n]+)",
        "title": r"(?:^|\n)Tytuł:\s*\n?([^\n]+)",
        "area": r"(?:^|\n)Obszar:\s*\n?([^\n]+)",
    }
    for key, pattern in fallback_patterns.items():
        if fields[key]:
            continue
        match = re.search(pattern, page_text, re.I)
        if match:
            fields[key] = _clean(match.group(1))
    msz_deadline_match = re.search(
        r"Termin przesłania oferty bliźniaczej do MSZ\s*:\s*([^\n]+)", page_text, re.I
    )
    beneficiary_deadline_match = re.search(
        r"Termin przesłania oferty przez MSZ do beneficjenta\s*:\s*([^\n]+)", page_text, re.I
    )
    fields.update(
        {
            "attachmentUrl": parser.attachment_url or None,
            "mszDeadline": _clean(msz_deadline_match.group(1)) if msz_deadline_match else None,
            "beneficiaryDeadline": _clean(beneficiary_deadline_match.group(1)) if beneficiary_deadline_match else None,
            "pageText": page_text,
            "sourceUrl": source_url,
        }
    )
    return fields


def _attachment_metadata(url: str | None) -> dict[str, str]:
    if not url:
        return {}
    try:
        _, headers, final_url = _request(url, timeout=20, method="HEAD")
    except TwinningError:
        return {"url": url}
    return {
        "url": final_url,
        "etag": headers.get("etag", ""),
        "lastModified": headers.get("last-modified", ""),
        "contentLength": headers.get("content-length", ""),
    }


def _offer_hash(payload: dict) -> str:
    stable = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(stable.encode("utf-8")).hexdigest()


def _parse_iso_date(value: str) -> date | None:
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except (TypeError, ValueError):
        return None


def list_active_offers(
    *,
    list_url: str = LIST_URL,
    lookback_days: int = 180,
    max_offers: int = 30,
    today: date | None = None,
) -> dict:
    current_date = today or date.today()
    listing = parse_listing(fetch_text(list_url))
    cutoff = current_date - timedelta(days=max(1, min(lookback_days, 730)))
    candidates = [offer for offer in listing if (_parse_iso_date(offer.published_date) or current_date) >= cutoff]
    candidates = candidates[: max(1, min(max_offers, 100))]

    active: list[dict] = []
    errors: list[dict] = []
    for offer in candidates:
        try:
            detail = parse_detail(fetch_text(offer.url), offer.url)
            deadline_date = parse_polish_date(detail.get("mszDeadline"))
            if deadline_date and deadline_date < current_date:
                continue
            attachment_meta = _attachment_metadata(detail.get("attachmentUrl"))
            payload = {
                "publishedDate": offer.published_date,
                "country": detail.get("beneficiary") or offer.country,
                "area": detail.get("area") or offer.area,
                "title": detail.get("title") or offer.title,
                "reference": detail.get("reference"),
                "url": offer.url,
                "attachmentUrl": detail.get("attachmentUrl"),
                "mszDeadline": detail.get("mszDeadline"),
                "beneficiaryDeadline": detail.get("beneficiaryDeadline"),
                "pageText": detail.get("pageText"),
                "attachmentMeta": attachment_meta,
            }
            payload["offerId"] = payload.get("reference") or offer.url.rstrip("/").rsplit("/", 1)[-1]
            payload["contentHash"] = _offer_hash(payload)
            active.append(payload)
        except Exception as exc:  # individual source errors must not hide other offers
            errors.append({"url": offer.url, "message": str(exc)})

    return {
        "source": list_url,
        "checkedAt": datetime.now().astimezone().isoformat(timespec="seconds"),
        "activeOffers": active,
        "errors": errors,
    }


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _document_score(name: str, text: str) -> int:
    normalized_name = name.lower()
    normalized = text.lower()
    score = 0
    if "annex c1" in normalized_name or "annex_c1" in normalized_name or "annexc1" in normalized_name:
        score += 100
    if "fiche" in normalized_name:
        score += 80
    markers = {
        "twinning fiche": 50,
        "profile and tasks of the project leader": 30,
        "profile and tasks of the rta": 30,
        "means/input from the eu member state": 30,
        "duration of the project": 15,
        "eu funded budget": 15,
    }
    for marker, weight in markers.items():
        if marker in normalized:
            score += weight
    return score


def _iter_documents(data: bytes, url: str) -> Iterable[tuple[str, bytes]]:
    lowered = url.lower().split("?", 1)[0]
    if lowered.endswith(".zip") or data[:4] == b"PK\x03\x04":
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            for info in archive.infolist():
                if info.is_dir() or info.file_size > MAX_DOCUMENT_BYTES:
                    continue
                suffix = info.filename.lower().rsplit(".", 1)[-1]
                if suffix not in {"pdf", "docx", "txt"}:
                    continue
                yield info.filename, archive.read(info)
        return
    name = url.rstrip("/").rsplit("/", 1)[-1] or "attachment"
    yield name, data


def _analysis_excerpt(text: str) -> str:
    cleaned = re.sub(r"[ \t]+", " ", text).strip()
    if len(cleaned) <= 25_000:
        return cleaned

    windows: list[tuple[int, int]] = [(0, 12_000)]
    markers = [
        "3.5 Components and results",
        "3.6 Means/input",
        "Profile and tasks of the Project Leader",
        "Profile and tasks of the RTA",
        "Profile and tasks of Component Leaders",
        "Profile and tasks of other short-term experts",
        "4. Budget",
        "6. Duration of the project",
        "12. Facilities available",
    ]
    lower = cleaned.lower()
    for marker in markers:
        # Fiches repeat section names in the table of contents. The last hit is
        # normally the substantive section with the requirements and tasks.
        index = lower.rfind(marker.lower())
        if index >= 0:
            windows.append((max(0, index - 700), min(len(cleaned), index + 5_500)))
    windows.sort()
    merged: list[tuple[int, int]] = []
    for start, end in windows:
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    excerpt = "\n\n[... kolejny istotny fragment ...]\n\n".join(cleaned[start:end] for start, end in merged)
    return excerpt[:MAX_ANALYSIS_CHARS]


def extract_offer(offer_url: str) -> dict:
    detail = parse_detail(fetch_text(offer_url), offer_url)
    attachment_url = detail.get("attachmentUrl")
    if not attachment_url:
        raise TwinningError("Strona oferty nie zawiera załącznika")
    data, _, final_url = _request(str(attachment_url), timeout=60)

    documents: list[dict] = []
    for name, content in _iter_documents(data, final_url):
        try:
            lowered = name.lower()
            if lowered.endswith(".pdf"):
                text = _extract_pdf(content)
            elif lowered.endswith(".docx"):
                text = _extract_docx(content)
            else:
                text = content.decode("utf-8", errors="replace")
            cleaned = _clean(text)
            if cleaned:
                documents.append(
                    {
                        "name": name,
                        "text": text,
                        "score": _document_score(name, text),
                        "characters": len(text),
                    }
                )
        except Exception as exc:
            documents.append({"name": name, "error": str(exc), "score": -1, "characters": 0})

    readable = [document for document in documents if document.get("text")]
    if not readable:
        raise TwinningError("Nie udało się odczytać żadnego dokumentu PDF/DOCX z załącznika")
    primary = max(readable, key=lambda document: (document["score"], document["characters"]))
    return {
        **detail,
        "attachmentUrl": final_url,
        "primaryDocument": primary["name"],
        "sourceFiles": [
            {key: value for key, value in document.items() if key != "text"}
            for document in documents
        ],
        "analysisText": _analysis_excerpt(str(primary["text"])),
        "analysisCharacters": len(str(primary["text"])),
    }
